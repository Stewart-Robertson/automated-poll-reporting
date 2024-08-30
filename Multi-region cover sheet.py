import pandas as pd
import docx
import pprint
import warnings
from normalise import clean_spaces
from normalise import standardNormalisation
from os import listdir
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
warnings.filterwarnings('ignore')

previous_files = [] # List of previous set of results manually pasted into here
# List of files to be normalised and then used to create the cover sheet
normed_files = []
for file in listdir('/path/to/my/directory'):
    normed_files.append(standardNormalisation())

normed_files = sorted(normed_files)
previous_files = sorted(previous_files)

# Function for bolding percentages
def pcs_bolding(table, col_idx, pcs_list):
    left_combined = round(pcs_list[0] + pcs_list[1], 2)
    right_combined = round(pcs_list[3] + pcs_list[4], 2)
    combined_pcs_list = [left_combined, pcs_list[2], right_combined, pcs_list[5]]
    max_val = max(combined_pcs_list)

    if left_combined == max_val:
        table.cell(2, col_idx).paragraphs[0].runs[0].bold = True
        table.cell(3, col_idx).paragraphs[0].runs[0].bold = True
    if right_combined == max_val:
        table.cell(5, col_idx).paragraphs[0].runs[0].bold = True
        table.cell(6, col_idx).paragraphs[0].runs[0].bold = True
    if pcs_list[2] == max_val:
        table.cell(4, col_idx).paragraphs[0].runs[0].bold = True
    if pcs_list[5] == max_val:
        table.cell(7, col_idx).paragraphs[0].runs[0].bold = True

# Function to add net figures to questions with "neither" that are to be combined
# Adds change in net figure from previous poll
def net_figure(table, col_idx, pcs_list_current, filename, file_list, skip, pcs_list_previous=None):
    if filename == file_list[0]:
        blank_row = table.add_row()
        net_row = table.add_row()
    # If skip == 0, the question exists in the previous file and so change should be included
    if skip == 0:
        if filename == file_list[0]:
            blank_row2 = table.add_row()
            change_row = table.add_row()
        left_combined_current = round(pcs_list_current[0] + pcs_list_current[1], 2)
        right_combined_current = round(pcs_list_current[3] + pcs_list_current[4], 2)
        left_combined_previous = round(pcs_list_previous[0] + pcs_list_previous[1], 2)
        right_combined_previous = round(pcs_list_previous[3] + pcs_list_previous[4], 2)
        # Calculate net percentages and get the change in net
        net_pc_current = left_combined_current - right_combined_current
        net_pc_previous = left_combined_previous - right_combined_previous
        change_pc = net_pc_current - net_pc_previous
        net_pc_current = "{:.0f}%".format((net_pc_current)*100)
        change_pc = "{:.0f}%".format((change_pc)*100)
        if not net_pc_current.startswith("-") and net_pc_current != "0%":
            net_pc_current = "+" + net_pc_current
        if not change_pc.startswith("-") and change_pc != "0%":
            change_pc = "+" + change_pc
        table.cell(len(table.rows) - 3, col_idx).text = net_pc_current
        table.cell(len(table.rows) - 3, col_idx).paragraphs[0].runs[0].italic = True
        table.cell(len(table.rows) - 1, col_idx).text = change_pc
        table.cell(len(table.rows) - 1, col_idx).paragraphs[0].runs[0].italic = True
        table.cell(len(table.rows) - 3, 0).text = "Net Result"
        table.cell(len(table.rows) - 3, 0).paragraphs[0].runs[0].italic = True
        table.cell(len(table.rows) - 1, 0).text = "Change from previous poll"
        table.cell(len(table.rows) - 1, 0).paragraphs[0].runs[0].italic = True
    # Combine answer codes two indices either side of "neither" for current and past files
    elif skip == 1:
        left_combined = round(pcs_list_current[0] + pcs_list_current[1], 2)
        right_combined = round(pcs_list_current[3] + pcs_list_current[4], 2)
        net_pc = "{:.0f}%".format((left_combined - right_combined)*100)
        if not net_pc.startswith("-") and net_pc != "0%":
            net_pc = "+" + net_pc
        table.cell(len(table.rows) - 1, col_idx).text = net_pc
        table.cell(len(table.rows) - 1, 0).text = "Net Result"
        table.cell(len(table.rows) - 1, col_idx).paragraphs[0].runs[0].italic = True
        table.cell(len(table.rows) - 1, 0).paragraphs[0].runs[0].italic = True

def combined_cover_sheet(list_of_files_current, list_of_files_previous):
    # Create document and set styling
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Gotham Book'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = Pt(0)

    print(list_of_files_current)

    # Read the first Excel file to get unique questions and attributes
    temp_df = pd.read_excel(list_of_files_current[0])
    qs = list(temp_df['Question'].unique())
    unique_attrs = temp_df['Attribute'].unique()
    pprint.pprint(unique_attrs)
    chosen_attr = str(input("Choose an Attribute: "))
    while chosen_attr not in unique_attrs:
        print("Please choose an attribute from the list above")
        chosen_attr = str(input("Choose an Attribute: "))

    # Initialise a dictionary to store the tables for each question
    # Each table can be accessed for matching questions in subsequent files using the question as a key
    question_tables = {}
    # Store answer codes with percentages either side of chosen answers with "neither" that are to be combined
    ans_to_combine = []
    # Standard questions will have their singular maximum percentage bolded
    ans_not_combine = []

    region_dfs = [pd.read_excel(file) for file in list_of_files_current]
    region_dfs_prev = [pd.read_excel(file) for file in list_of_files_previous]

    # Iterate through the unique questions in the current files to create tables
    for q in qs:
        print(q)
        sample_sizes = []
        for df in region_dfs:
            df_q_attr = df[(df['Question'] == q) & (df['Attribute'] == chosen_attr)]
            sample_size = df_q_attr['Unweighted Total'].unique()[0]
            sample_sizes.append(sample_size)
        para = doc.add_paragraph(q)
        para.style = doc.styles['Heading 2']
        doc.add_paragraph(" ")

        table = doc.add_table(rows=1, cols=len(list_of_files_current) + 1)
        table.columns[0].width= Inches(12)
        question_tables[q] = table

        heading_row = table.rows[0]
        for idx, filename in enumerate(list_of_files_current):
            heading_row.cells[idx + 1].text = filename + " (n=" + str(sample_sizes[idx]) + ")"
        table.add_row()
        doc.add_paragraph(" ")

    # Loop through all the regions and add their data to their matching table
    for filename, df in zip(list_of_files_current, region_dfs):
        df = df[df['Attribute'] == chosen_attr].reset_index()

        # If the region has a poll in the previous files list, read that in as a Dataframe
        for fle in list_of_files_previous:
            if fle.split("US")[1] in filename:
                df_prev = pd.read_excel(fle)
                df_prev = df_prev[df_prev['Attribute'] == chosen_attr].reset_index()
            else:
                continue
                
        for q in qs:
            ans_per_q = df[df["Question"] == q]
            answers = list(ans_per_q['Answer'].unique())
            pcs_per_q = list(ans_per_q['% Votes'])
            pcs_per_q = [round(pcs, 2) for pcs in pcs_per_q]
            table = question_tables[q]

            # Get the same for the previous poll (don't need answers) if the question in the current file is in the previous file
            if q in df_prev['Question'].values:
                skip = 0
                ans_per_q_prev = df_prev[df_prev["Question"] == q]
                pcs_per_q_prev = list(ans_per_q_prev['% Votes'])
                pcs_per_q_prev = [round(pcs, 2) for pcs in pcs_per_q_prev]
            else:
                skip = 1
                pcs_per_q_prev = None
                pass

            # If it's the first file, add the answers to the table in column 0
            if filename == list_of_files_current[0]:
                for answer in answers:
                    new_row = table.add_row()
                    new_row.cells[0].text = answer

            # Add the percentages for each answer in the current file's column        
            for idx, answer in enumerate(answers):
                col_idx = list_of_files_current.index(filename) + 1
                pctg = ans_per_q[ans_per_q['Answer'] == answer]['% Votes'].values[0]
                pctg_text = "{:.0f}%".format(pctg * 100)
                answer_row_idx = [row.cells[0].text for row in table.rows[1:]].index(answer) + 1
                table.cell(answer_row_idx, col_idx).text = pctg_text
            
            # Initial check to see if answers either side of "neither" answer codes are in the list of answers to be combined
            if answers in ans_to_combine:
                pcs_bolding(table, col_idx, pcs_per_q)
                net_figure(table=table, col_idx=col_idx, pcs_list_current=pcs_per_q, filename=filename, file_list=list_of_files_current, skip=skip, pcs_list_previous=pcs_per_q_prev)
            
            # If the answers are explicitly not to be combined, bold only the highest % in the column
            elif answers in ans_not_combine:
                bold_idx = [i for i, x in enumerate(pcs_per_q) if x == max(pcs_per_q)]
                if len(bold_idx) == 1:
                    table.cell(bold_idx[0] + 2, col_idx).paragraphs[0].runs[0].bold = True
                else:
                    for i in bold_idx:
                        table.cell(i + 2, col_idx).paragraphs[0].runs[0].bold = True

            # Check if the first instance of questions containing "neither" answer codes are to be combined
            elif answers not in ans_to_combine and "neither" in " ".join(answers).lower():
                print(answers)
                comb = str(input("Would you like to combine these answers? (y/n) "))
                if comb.lower() == "y":
                    ans_to_combine.append(answers)
                    pcs_bolding(table, col_idx, pcs_per_q)
                    net_figure(table=table, col_idx=col_idx, pcs_list_current=pcs_per_q, filename=filename, file_list=list_of_files_current, skip=skip, pcs_list_previous=pcs_per_q_prev)
                elif comb.lower() == "n":
                    ans_not_combine.append(answers)
                    bold_idx = [i for i, x in enumerate(pcs_per_q) if x == max(pcs_per_q)]
                    if len(bold_idx) == 1:
                        table.cell(bold_idx[0] + 2, col_idx).paragraphs[0].runs[0].bold = True
                    else:
                        for i in bold_idx:
                            table.cell(i + 2, col_idx).paragraphs[0].runs[0].bold = True
                else:
                    print("Please enter 'y' or 'n'")
            # If "n" to the above question, bold the column according to the highest % and add it to answers explicitly not to be combined
            else:
                bold_idx = [i for i, x in enumerate(pcs_per_q) if x == max(pcs_per_q)]
                if len(bold_idx) == 1:
                    table.cell(bold_idx[0] + 2, col_idx).paragraphs[0].runs[0].bold = True
                else:
                    for i in bold_idx:
                        table.cell(i + 2, col_idx).paragraphs[0].runs[0].bold = True

    doc.save('Combined Cover Sheet.docx')
    print("Combined cover sheet created")

if __name__ == "__main__":
    combined_cover_sheet(normed_files, previous_files)
